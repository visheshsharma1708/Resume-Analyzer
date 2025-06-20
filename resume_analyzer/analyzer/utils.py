from django.shortcuts import render
from .forms import ResumeUploadForm
from .models import Candidate
import PyPDF2
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

def extract_text(file):
    if file.name.endswith('.pdf'):
        reader = PyPDF2.PdfReader(file)
        return " ".join([page.extract_text() or "" for page in reader.pages])
    elif file.name.endswith('.docx'):
        doc = docx.Document(file)
        return " ".join([para.text for para in doc.paragraphs])
    return ""

def get_match_score(resume_text, job_desc):
    vectorizer = TfidfVectorizer(stop_words='english')
    vectors = vectorizer.fit_transform([resume_text, job_desc])
    score = cosine_similarity(vectors[0:1], vectors[1:2])[0][0]
    resume_words = set(resume_text.lower().split())
    job_words = set(job_desc.lower().split())
    missing = list(job_words - resume_words)
    return round(score * 100, 2), missing[:10]

def extract_text_from_resume(resume_file):
    text = ""
    if resume_file.name.endswith('.pdf'):
        reader = PyPDF2.PdfReader(resume_file)
        for page in reader.pages:
            text += page.extract_text() or ""
    elif resume_file.name.endswith('.docx'):
        doc = docx.Document(resume_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    return text

def calculate_match_score(resume_text, job_description):
    vectorizer = TfidfVectorizer(stop_words='english')
    vectors = vectorizer.fit_transform([resume_text, job_description])
    score = cosine_similarity(vectors[0:1], vectors[1:2])[0][0]
    resume_words = set(resume_text.lower().split())
    job_words = set(job_description.lower().split())
    missing = list(job_words - resume_words)
    return round(score * 100, 2), missing[:10]

def upload_resume(request):
    if request.method == "POST":
        form = ResumeUploadForm(request.POST, request.FILES)
        if form.is_valid():
            candidate = form.save()
            resume_text = extract_text_from_resume(candidate.resume)
            job_description = request.POST.get("job_description", "")
            match_score, missing_keywords = calculate_match_score(resume_text, job_description)
            return render(request, 'result.html', {
                'candidate': candidate,
                'match_score': match_score,
                'missing_keywords': missing_keywords
            })
    else:
        form = ResumeUploadForm()
    return render(request, 'upload.html', {'form': form})
